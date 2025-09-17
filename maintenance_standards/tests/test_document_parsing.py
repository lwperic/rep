import os
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from backend.core.document_manager.parser import DocumentParser
from backend.models.document_structure import ParagraphType

def create_test_doc_with_structure(path: str):
    """创建包含结构化内容的测试文档"""
    doc = Document()
    
    # 确保文档有Title样式
    styles = doc.styles
    if 'Title' not in styles:
        styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    
    # 添加文档标题（使用Title样式）
    title = doc.add_paragraph("维修标准文档")
    title.style = doc.styles['Title']
    
    # 添加第一章（使用Heading 1样式）
    heading1 = doc.add_paragraph("第一章：概述")
    heading1.style = doc.styles['Heading 1']
    
    # 添加1.1节（使用Heading 2样式）
    heading2 = doc.add_paragraph("1.1 文档目的")
    heading2.style = doc.styles['Heading 2']
    doc.add_paragraph("本文档旨在规范维修流程。")
    
    # 添加1.2节（使用Heading 2样式）
    heading2 = doc.add_paragraph("1.2 适用范围")
    heading2.style = doc.styles['Heading 2']
    
    # 添加列表项
    doc.add_paragraph("● 机械设备维修")
    doc.add_paragraph("● 电气设备维修")
    
    # 设置表格标题
    caption = doc.add_paragraph("表1-1 设备维修周期")
    caption.style = 'Caption'
    
    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    
    # 填充表格内容
    cells = table.rows[0].cells
    cells[0].text = "设备类型"
    cells[1].text = "维修周期"
    cells = table.rows[1].cells
    cells[0].text = "机械设备"
    cells[1].text = "每季度"
    
    # 添加引用
    doc.add_paragraph('"设备维修应当遵循安全第一的原则。"')
    
    doc.save(path)

def test_document_parser():
    """测试文档解析器功能"""
    test_doc_path = "test_parsing.docx"
    
    try:
        # 创建测试文档
        create_test_doc_with_structure(test_doc_path)
        
        # 解析文档
        parser = DocumentParser()
        structure = parser.parse_document(test_doc_path)
        
        # 验证段落
        assert len(structure.paragraphs) > 0
        
        # 查找不同类型的段落
        titles = [p for p in structure.paragraphs if p.type == ParagraphType.TITLE]
        list_items = [p for p in structure.paragraphs if p.type == ParagraphType.LIST_ITEM]
        references = [p for p in structure.paragraphs if p.type == ParagraphType.REFERENCE]
        
        # 验证文档标题
        assert structure.title == "维修标准文档"
        
        # 验证标题层级
        assert len(titles) >= 3  # 文档标题 + 第一章 + 两个子节
        chapter_title = next(t for t in titles if "第一章" in t.text)
        assert chapter_title.level == 1  # 第一章应该是1级标题
        
        # 验证列表项
        assert len(list_items) == 2
        assert any("机械设备" in item.text for item in list_items)
        assert any("电气设备" in item.text for item in list_items)
        
        # 验证引用
        assert len(references) == 1
        assert "安全第一" in references[0].text
        
        # 验证章节结构
        assert len(structure.sections) > 0
        first_section = structure.sections[0]
        assert first_section.level == 1  # 第一章的层级应该是1
        assert len(first_section.subsections) == 2  # 应该有两个子节
        
        # 验证表格
        assert len(structure.tables) > 0
        first_table = structure.tables[0]
        assert first_table.num_rows == 2
        assert first_table.num_cols == 2
        assert first_table.caption == "表1-1 设备维修周期"  # 验证表格标题
        
        # 打印解析结果
        print("\n解析结果:")
        print(f"文档标题: {structure.title}")
        print(f"总段落数: {len(structure.paragraphs)}")
        print(f"标题数量: {len(titles)}")
        print(f"列表项数量: {len(list_items)}")
        print(f"引用数量: {len(references)}")
        print(f"章节数: {len(structure.sections)}")
        print(f"表格数: {len(structure.tables)}")
        
    finally:
        # 清理测试文件
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
