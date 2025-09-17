import os
import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from backend.core.document_manager.parser import DocumentParser
from backend.models.document_structure import ParagraphType

def create_doc_with_complex_tables(path: str):
    """创建包含复杂表格的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("复杂表格测试", 0)
    
    # 添加带有合并单元格的表格
    doc.add_paragraph("以下是一个包含合并单元格的表格：")
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    
    # 填充表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "维修项目"
    header_cells[1].text = "检查周期"
    header_cells[2].text = "维修要求"
    header_cells[3].text = "备注"
    
    # 合并单元格示例
    cell1 = table.cell(1, 0)
    cell2 = table.cell(2, 0)
    cell1.merge(cell2)
    cell1.text = "机械传动系统"
    
    # 添加表格标题
    caption = doc.add_paragraph("表2-1 设备维修计划")
    caption.style = 'Caption'
    
    doc.save(path)

def create_doc_with_complex_lists(path: str):
    """创建包含复杂列表的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("复杂列表测试", 0)
    
    # 添加多级列表
    doc.add_paragraph("1. 维修准备", style='List Number')
    doc.add_paragraph("   a) 工具准备", style='List Bullet 2')
    doc.add_paragraph("   b) 安全防护", style='List Bullet 2')
    doc.add_paragraph("2. 维修步骤", style='List Number')
    doc.add_paragraph("   a) 检查部件", style='List Bullet 2')
    doc.add_paragraph("   b) 更换零件", style='List Bullet 2')
    
    doc.save(path)

def create_doc_with_references(path: str):
    """创建包含多种引用的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading("引用测试", 0)
    
    # 添加交叉引用
    doc.add_paragraph("参考第一章第二节的维修流程...")
    doc.add_paragraph("如图1-2所示的装配图...")
    doc.add_paragraph("根据表2-1的维修计划...")
    
    # 添加引用文本
    doc.add_paragraph('"设备维修必须按照标准流程进行。"')
    doc.add_paragraph('"确保安全是首要任务。"')
    
    doc.save(path)

def test_complex_tables():
    """测试复杂表格解析"""
    test_doc_path = "test_complex_tables.docx"
    try:
        create_doc_with_complex_tables(test_doc_path)
        parser = DocumentParser()
        structure = parser.parse_document(test_doc_path)
        
        # 验证表格解析
        assert len(structure.tables) > 0
        table = structure.tables[0]
        
        # 验证表格标题
        assert table.caption == "表2-1 设备维修计划"
        
        # 验证表格维度
        assert table.num_rows == 4
        assert table.num_cols == 4
        
        # 验证合并单元格
        merged_cells = [cell for cell in table.cells if cell.text == "机械传动系统"]
        assert len(merged_cells) == 1
        
    finally:
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)

def test_complex_lists():
    """测试复杂列表解析"""
    test_doc_path = "test_complex_lists.docx"
    try:
        create_doc_with_complex_lists(test_doc_path)
        parser = DocumentParser()
        structure = parser.parse_document(test_doc_path)
        
        # 获取所有列表项
        list_items = [p for p in structure.paragraphs if p.type == ParagraphType.LIST_ITEM]
        
        # 验证列表项数量
        assert len(list_items) == 6
        
        # 验证列表层级
        first_level_items = [item for item in list_items if "1." in item.text or "2." in item.text]
        second_level_items = [item for item in list_items if "a)" in item.text or "b)" in item.text]
        
        assert len(first_level_items) == 2
        assert len(second_level_items) == 4
        
    finally:
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)

def test_references():
    """测试引用解析"""
    test_doc_path = "test_references.docx"
    try:
        create_doc_with_references(test_doc_path)
        parser = DocumentParser()
        structure = parser.parse_document(test_doc_path)
        
        # 获取所有引用
        references = [p for p in structure.paragraphs if p.type == ParagraphType.REFERENCE]
        
        # 验证引用数量
        assert len(references) == 2
        
        # 验证引用内容
        assert any("标准流程" in ref.text for ref in references)
        assert any("安全是首要任务" in ref.text for ref in references)
        
    finally:
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)

if __name__ == "__main__":
    pytest.main([__file__, "-v"])