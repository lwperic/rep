from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from backend.core.document_manager.parser import DocumentParser
from pathlib import Path

def create_demo_document(path: str):
    """创建一个包含各种元素的示例维修标准文档"""
    doc = Document()
    
    # 添加样式
    styles = doc.styles
    if 'Title' not in styles:
        styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    
    # 文档标题
    title = doc.add_paragraph("发动机维修标准手册")
    title.style = doc.styles['Title']
    
    # 第一章：概述
    heading1 = doc.add_paragraph("第一章：概述")
    heading1.style = doc.styles['Heading 1']
    
    # 1.1 目的
    heading2 = doc.add_paragraph("1.1 目的")
    heading2.style = doc.styles['Heading 2']
    doc.add_paragraph("本手册旨在规范发动机维修流程，确保维修质量和安全。")
    
    # 1.2 适用范围
    heading2 = doc.add_paragraph("1.2 适用范围")
    heading2.style = doc.styles['Heading 2']
    doc.add_paragraph("● 汽油发动机维修")
    doc.add_paragraph("● 柴油发动机维修")
    
    # 第二章：维修准备
    heading1 = doc.add_paragraph("第二章：维修准备")
    heading1.style = doc.styles['Heading 1']
    
    # 2.1 工具准备
    heading2 = doc.add_paragraph("2.1 工具准备")
    heading2.style = doc.styles['Heading 2']
    
    # 添加表格
    doc.add_paragraph("维修所需工具清单如下：")
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 设置表格标题
    caption = doc.add_paragraph("表2-1 维修工具清单")
    caption.style = 'Caption'
    
    # 填充表格
    header_cells = table.rows[0].cells
    header_cells[0].text = "工具名称"
    header_cells[1].text = "规格"
    header_cells[2].text = "数量"
    
    data = [
        ["扳手", "8-19mm", "1套"],
        ["螺丝刀", "十字/一字", "各1把"],
        ["扭力扳手", "20-200N·m", "1把"]
    ]
    
    for i, row in enumerate(data):
        cells = table.rows[i + 1].cells
        for j, value in enumerate(row):
            cells[j].text = value
    
    # 2.2 安全事项
    heading2 = doc.add_paragraph("2.2 安全事项")
    heading2.style = doc.styles['Heading 2']
    doc.add_paragraph('"进行维修工作时必须遵守安全规程，确保人身和设备安全。"')
    
    # 添加图片（这里使用空白图片作为示例）
    doc.add_picture('placeholder.png', width=Inches(4))
    # 添加图片标题
    img_caption = doc.add_paragraph("图2-1 安全防护装备示意图")
    img_caption.style = 'Caption'
    
    doc.save(path)

def main():
    # 创建演示文档
    demo_doc_path = "维修标准示例.docx"
    create_demo_document(demo_doc_path)
    
    # 解析文档
    parser = DocumentParser()
    structure = parser.parse_document(demo_doc_path)
    
    # 打印解析结果
    print("\n=== 文档解析结果 ===")
    print(f"文档标题: {structure.title}")
    print(f"\n=== 章节结构 ===")
    for section in structure.sections:
        print(f"\n{section.title.text} (Level {section.level})")
        for subsection in section.subsections:
            print(f"  |- {subsection.title.text}")
            for para in subsection.paragraphs:
                print(f"     |- {para.text[:50]}..." if len(para.text) > 50 else f"     |- {para.text}")
    
    print(f"\n=== 表格信息 ===")
    for i, table in enumerate(structure.tables, 1):
        print(f"\n表格 {i}:")
        print(f"标题: {table.caption}")
        print(f"大小: {table.num_rows}行 x {table.num_cols}列")
        print("内容预览:")
        cells_by_row = {}
        for cell in table.cells:
            if cell.row not in cells_by_row:
                cells_by_row[cell.row] = []
            cells_by_row[cell.row].append(cell.text)
        
        for row in sorted(cells_by_row.keys()):
            print(f"  {' | '.join(cells_by_row[row])}")
    
    print(f"\n=== 图片信息 ===")
    for i, image in enumerate(structure.images, 1):
        print(f"\n图片 {i}:")
        print(f"路径: {image.path}")
        print(f"标题: {image.caption}")

if __name__ == "__main__":
    main()