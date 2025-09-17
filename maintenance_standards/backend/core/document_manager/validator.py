import os
import tempfile
from pathlib import Path
from typing import Optional
from loguru import logger
from docx import Document as DocxDocument

from ...config import settings
from ...utils.error_handler import DocumentError

class DocumentValidator:
    """文档验证器"""
    
    def validate_file(self, filename: str, file_content: bytes) -> None:
        """验证文件
        
        Args:
            filename: 文件名
            file_content: 文件内容
            
        Raises:
            DocumentError: 当验证失败时
        """
        # 验证文件扩展名
        self._validate_extension(filename)
        
        # 验证文件大小
        self._validate_file_size(file_content)
        
        # 验证文件格式
        self._validate_file_format(file_content)
        
        # 验证文件内容
        self._validate_content(file_content)
    
    def _validate_extension(self, filename: str) -> None:
        """验证文件扩展名"""
        ext = Path(filename).suffix.lower()
        if ext not in settings.ALLOWED_EXTENSIONS:
            raise DocumentError(f"不支持的文件类型: {ext}")
    
    def _validate_file_size(self, file_content: bytes) -> None:
        """验证文件大小"""
        if len(file_content) > settings.MAX_FILE_SIZE:
            size_mb = len(file_content) / (1024 * 1024)
            max_size_mb = settings.MAX_FILE_SIZE / (1024 * 1024)
            raise DocumentError(
                f"文件大小超过限制: {size_mb:.1f}MB (最大允许 {max_size_mb:.1f}MB)"
            )
    
    def _validate_file_format(self, file_content: bytes) -> None:
        """验证文件格式"""
        try:
            # 将文件内容写入临时文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file.flush()
                
                # 尝试打开文档
                _ = DocxDocument(tmp_file.name)
                
        except Exception as e:
            raise DocumentError(f"文件格式无效: {str(e)}")
        finally:
            if 'tmp_file' in locals():
                os.unlink(tmp_file.name)
    
    def validate_content_for_extraction(self, file_path: str) -> bool:
        """验证文档内容是否适合进行知识图谱提取
        
        Args:
            file_path: 文档路径
            
        Returns:
            bool: 文档内容是否有效
        """
        try:
            doc = DocxDocument(file_path)
            
            # 验证文档是否为空
            if len(doc.paragraphs) == 0:
                logger.warning("文档内容为空")
                return False
                
            # 检查必需的内容段落
            required_sections = ["步骤", "工具", "安全", "注意"]
            found_sections = set()
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # 检查段落是否包含必需的关键词
                for section in required_sections:
                    if section in text:
                        found_sections.add(section)
                        
            # 检查是否找到所有必需的部分
            missing_sections = set(required_sections) - found_sections
            if missing_sections:
                logger.warning(f"缺少必需的内容部分: {', '.join(missing_sections)}")
                return False
                
            # 验证段落结构
            current_section = None
            section_content = {}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                # 判断是否是新的段落标题
                is_title = any(text.startswith(section) for section in required_sections)
                
                if is_title:
                    current_section = next(s for s in required_sections if text.startswith(s))
                    section_content[current_section] = []
                elif current_section:
                    section_content[current_section].append(text)
                    
            # 检查每个部分是否有内容
            for section in required_sections:
                if section not in section_content or not section_content[section]:
                    logger.warning(f"部分 '{section}' 没有具体内容")
                    return False
                    
            return True
            
        except Exception as e:
            logger.error(f"文档内容验证失败: {str(e)}")
            return False
        """验证文件格式"""
        try:
            # 创建临时文件进行格式验证
            temp_path = Path(settings.UPLOAD_FOLDER) / ".temp.docx"
            with open(temp_path, "wb") as f:
                f.write(file_content)
            
            # 尝试打开文档
            DocxDocument(temp_path)
            
            # 删除临时文件
            os.remove(temp_path)
        except Exception as e:
            logger.error(f"Invalid document format: {str(e)}")
            raise DocumentError("文档格式无效或已损坏")
    
    def _validate_content(self, file_content: bytes) -> None:
        """验证文件内容"""
        try:
            # 创建临时文件进行内容验证
            temp_path = Path(settings.UPLOAD_FOLDER) / ".temp.docx"
            with open(temp_path, "wb") as f:
                f.write(file_content)
            
            doc = DocxDocument(temp_path)
            
            # 验证文档是否为空
            if len(doc.paragraphs) == 0:
                raise DocumentError("文档内容为空")
            
            # 验证文档结构
            self._validate_document_structure(doc)
            
            # 删除临时文件
            os.remove(temp_path)
        except DocumentError:
            raise
        except Exception as e:
            logger.error(f"Content validation error: {str(e)}")
            raise DocumentError("文档内容验证失败")
    
    def _validate_document_structure(self, doc: DocxDocument) -> None:
        """验证文档结构"""
        # TODO: 实现更详细的文档结构验证逻辑
        # 例如：验证标题层级、必要章节等
        pass