from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re
from docx import Document
from loguru import logger

class DocumentCleaner:
    """文档清洗器，用于清洗和标准化文档内容"""
    
    def __init__(self):
        # 清洗规则配置
        self.rules = {
            "remove_extra_spaces": True,  # 移除多余空格
            "normalize_punctuation": True,  # 标准化标点符号
            "fix_line_breaks": True,      # 修复换行
            "remove_empty_paragraphs": True,  # 移除空段落
        }
    
    def clean_text(self, text: str) -> str:
        """清洗文本内容
        
        Args:
            text: 待清洗的文本
            
        Returns:
            清洗后的文本
        """
        if not text:
            return text
            
        # 移除多余空格
        if self.rules["remove_extra_spaces"]:
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
        
        # 标准化标点符号
        if self.rules["normalize_punctuation"]:
            text = self._normalize_punctuation(text)
        
        return text
    
    def clean_document(self, doc_path: str) -> Tuple[Document, Dict]:
        """清洗文档内容
        
        Args:
            doc_path: 文档路径
            
        Returns:
            清洗后的文档对象和清洗统计信息
        """
        try:
            # 打开文档
            doc = Document(doc_path)
            stats = {
                "total_paragraphs": len(doc.paragraphs),
                "cleaned_paragraphs": 0,
                "removed_paragraphs": 0
            }
            
            # 清洗段落
            paragraphs_to_remove = []
            for i, para in enumerate(doc.paragraphs):
                if not para.text.strip() and self.rules["remove_empty_paragraphs"]:
                    paragraphs_to_remove.append(i)
                    stats["removed_paragraphs"] += 1
                    continue
                
                # 清洗文本
                original_text = para.text
                cleaned_text = self.clean_text(original_text)
                
                if cleaned_text != original_text:
                    para.text = cleaned_text
                    stats["cleaned_paragraphs"] += 1
            
            # 移除空段落（从后向前移除，避免索引变化）
            for i in reversed(paragraphs_to_remove):
                doc._body._body.remove(doc._body._body[i])
            
            logger.info(f"Document cleaned: {stats}")
            return doc, stats
            
        except Exception as e:
            logger.error(f"Error cleaning document: {str(e)}")
            raise
    
    def _normalize_punctuation(self, text: str) -> str:
        """标准化标点符号
        
        Args:
            text: 待处理的文本
            
        Returns:
            处理后的文本
        """
        # 中文标点符号标准化
        punctuation_map = {
            '，': ',',  # 全角逗号转半角
            '。': '.',  # 句号转点号
            '：': ':',  # 全角冒号转半角
            '；': ';',  # 全角分号转半角
            '"': '"',   # 全角引号转半角
            '"': '"',   # 全角引号转半角
            ''': "'",   # 全角单引号转半角
            ''': "'",   # 全角单引号转半角
            '！': '!',  # 全角感叹号转半角
            '？': '?',  # 全角问号转半角
            '（': '(',  # 全角括号转半角
            '）': ')',  # 全角括号转半角
            '【': '[',  # 全角方括号转半角
            '】': ']',  # 全角方括号转半角
            '《': '<',  # 全角尖括号转半角
            '》': '>',  # 全角尖括号转半角
        }
        
        for old, new in punctuation_map.items():
            text = text.replace(old, new)
        
        # 处理重复的标点符号
        text = re.sub(r'([,.!?;])\1+', r'\1', text)
        
        # 确保标点符号后有空格
        text = re.sub(r'([,.!?;:])([^\s])', r'\1 \2', text)
        
        return text