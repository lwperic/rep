import re
from typing import List, Tuple, Dict, Optional
from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.table import Table as DocxTable
from pathlib import Path
from loguru import logger

from ...models.document_structure import (
    DocumentStructure,
    Section,
    Paragraph,
    Table,
    TableCell,
    Image,
    ParagraphType
)

class DocumentParser:
    """文档解析器，用于解析文档结构和内容"""
    
    def __init__(self):
        # 标题模式配置
        self.title_patterns = {
            r"^第[一二三四五六七八九十]+章\s*[:：]?\s*(.*)$": 1,  # 章
            r"^\d+\.\s+(.*)$": 1,  # 1. 标题
            r"^\d+\.\d+\s+(.*)$": 2,  # 1.1 标题
            r"^\d+\.\d+\.\d+\s+(.*)$": 3,  # 1.1.1 标题
        }
        
        # 列表项模式
        self.list_patterns = [
            r"^[●○•]\s+.*$",  # 实心圆点、空心圆点
            r"^\(?[0-9a-zA-Z]\)?[\s.、].*$",  # (1) 或 1. 或 1、
            r"^[-—]\s+.*$",  # 短横线
        ]
    
    def parse_document(self, doc_path: str) -> DocumentStructure:
        """解析文档
        
        Args:
            doc_path: 文档路径
            
        Returns:
            解析后的文档结构
        """
        try:
            # 打开文档
            doc = Document(doc_path)
            
            # 初始化文档结构
            structure = DocumentStructure()
            
            # 解析段落
            self._parse_paragraphs(doc, structure)
            
            # 解析文档结构（章节）
            self._parse_sections(structure)
            
            # 解析表格
            self._parse_tables(doc, structure)
            
            # 解析图片
            self._parse_images(doc, structure)
            
            logger.info(f"Document parsed successfully: {doc_path}")
            return structure
            
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            raise
    
    def _parse_paragraphs(self, doc: _Document, structure: DocumentStructure):
        """解析段落"""
        for idx, para in enumerate(doc.paragraphs):
            # 跳过空段落
            if not para.text.strip():
                continue
            
            # 判断段落类型
            para_type = self._determine_paragraph_type(para)
            level = 0
            
            if para_type == ParagraphType.TITLE:
                level = self._get_title_level(para)
            
            # 创建段落对象
            paragraph = Paragraph(
                text=para.text,
                type=para_type,
                level=level,
                index=idx,
                style=para.style.name if para.style else None
            )
            
            # 添加到文档结构
            structure.paragraphs.append(paragraph)
            
            # 如果是第一个标题，可能是文档标题
            if idx == 0 and para_type == ParagraphType.TITLE:
                structure.title = para.text
    
    def _parse_sections(self, structure: DocumentStructure):
        """解析文档章节结构"""
        current_sections = []  # [(level, section), ...]
        
        for idx, para in enumerate(structure.paragraphs):
            if para.type != ParagraphType.TITLE:
                # 将非标题段落添加到当前最深层的章节
                if current_sections:
                    current_sections[-1][1].paragraphs.append(para)
                continue
            
            # 忽略文档标题样式的段落
            if para.style == 'Title':
                continue
                
            level = para.level
            
            # 创建新的章节（确保level至少为1）
            new_section = Section(
                title=para,
                level=max(1, level),  # 确保章节层级从1开始
                start_index=idx,
                end_index=idx  # 暂时设置为当前索引，后续会更新
            )
            
            # 根据层级关系处理章节
            while current_sections and current_sections[-1][0] >= level:
                current_sections[-1][1].end_index = idx - 1
                current_sections.pop()
            
            if current_sections:
                # 添加为当前章节的子章节
                current_sections[-1][1].subsections.append(new_section)
            else:
                # 添加为顶级章节
                structure.sections.append(new_section)
            
            current_sections.append((level, new_section))
    
    def _parse_tables(self, doc: _Document, structure: DocumentStructure):
        """解析表格"""
        table_caption = None
        table_pending = False
        
        # 遍历所有段落，提取表格标题和表格内容
        for element in doc._body._element:
            if element.tag.endswith('p'):
                # 处理段落，查找表格标题
                para = DocxParagraph(element, doc._parent)
                try:
                    if para.style and para.style.name == 'Caption':
                        table_caption = para.text.strip()
                        table_pending = True
                except AttributeError:
                    # 如果获取样式失败，尝试通过其他方式判断是否是标题
                    if para.text.strip().startswith('表'):
                        table_caption = para.text.strip()
                        table_pending = True
            elif element.tag.endswith('tbl'):
                # 处理表格
                table = DocxTable(element, doc._parent)
                parsed_table = self._parse_table(table, table_caption)
                if table_caption:
                    parsed_table.caption = table_caption
                structure.tables.append(parsed_table)
                # 重置标题状态
                table_caption = None
                table_pending = False
    
    def _parse_table(self, table: DocxTable, caption: Optional[str] = None) -> Table:
        """解析单个表格
        
        Args:
            table: 要解析的表格对象
            caption: 可选的表格标题
            
        Returns:
            解析后的表格对象
        """
        cells = []
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell_obj = TableCell(
                    text=cell.text.strip(),
                    row=i,
                    col=j,
                    is_header=(i == 0),  # 假设第一行是表头
                )
                cells.append(cell_obj)
        
        return Table(
            cells=cells,
            num_rows=num_rows,
            num_cols=num_cols,
            caption=caption
        )
    
    def _parse_images(self, doc: _Document, structure: DocumentStructure):
        """解析图片"""
        try:
            # 获取文档目录
            doc_dir = Path(doc._path).parent if hasattr(doc, '_path') else None
            
            # 存储图片序号计数
            figure_counter = 1
            # 存储当前章节编号
            current_chapter = 1
            # 存储待处理的图片标题
            image_caption = None
            image_pending = False
            
            # 遍历文档中的所有元素
            for element in doc._body._element:
                if element.tag.endswith('p'):
                    # 处理段落，查找图片标题
                    para = DocxParagraph(element, doc._parent)
                    try:
                        # 检查是否是新章节
                        if para.style and para.style.name.startswith('Heading 1'):
                            # 从文本中提取章节编号
                            match = re.match(r'^第([一二三四五六七八九十]+)章', para.text)
                            if match:
                                # 将中文数字转换为阿拉伯数字
                                chinese_nums = {'一':1, '二':2, '三':3, '四':4, '五':5,
                                              '六':6, '七':7, '八':8, '九':9, '十':10}
                                current_chapter = chinese_nums.get(match.group(1), current_chapter)
                                figure_counter = 1  # 重置图片计数器
                                
                        # 检查是否是图片标题
                        if (para.style and para.style.name == 'Caption' and 
                            para.text.strip().startswith('图')):
                            image_caption = para.text.strip()
                            image_pending = True
                        elif para.text.strip().startswith('图'):  # 备选方案
                            image_caption = para.text.strip()
                            image_pending = True
                    except AttributeError:
                        pass
                        
                elif element.tag.endswith(('drawing', 'pict')):
                    # 处理图片元素
                    for rel in doc.part.rels.values():
                        if "image" in rel.reltype:
                            # 构建图片路径
                            image_path = (doc_dir / rel._target) if doc_dir else rel._target
                            
                            if doc_dir is None or image_path.exists():
                                # 如果没有明确的标题，生成默认标题
                                if not image_caption:
                                    image_caption = f"图{current_chapter}-{figure_counter}"
                                
                                # 创建图片对象
                                image = Image(
                                    path=str(image_path),
                                    caption=image_caption,
                                    figure_number=f"{current_chapter}-{figure_counter}"
                                )
                                
                                # 更新计数器
                                figure_counter += 1
                                
                                # 添加到文档结构
                                structure.images.append(image)
                                
                                # 重置标题状态
                                image_caption = None
                                image_pending = False
                        
        except Exception as e:
            logger.warning(f"Error parsing images: {str(e)}")
            # 继续处理，不中断文档解析
    
    def _determine_paragraph_type(self, para: DocxParagraph) -> ParagraphType:
        """判断段落类型"""
        text = para.text.strip()
        
        # 安全地检查段落样式
        try:
            if para.style:
                style_name = para.style.name
                if style_name == 'Title' or style_name.startswith('Heading'):
                    return ParagraphType.TITLE
                elif style_name == 'Caption':
                    return ParagraphType.CONTENT  # 表格标题作为特殊内容处理
        except AttributeError:
            # 如果获取样式失败，继续使用其他方法判断
            pass
        
        # 检查是否是标题
        if any(re.match(pattern, text) for pattern in self.title_patterns.keys()):
            return ParagraphType.TITLE
        
        # 检查是否是列表项
        if any(re.match(pattern, text) for pattern in self.list_patterns):
            return ParagraphType.LIST_ITEM
        
        # 检查是否是引用
        if text.startswith('"') and text.endswith('"'):
            return ParagraphType.REFERENCE
        
        # 默认为正文
        return ParagraphType.CONTENT
    
    def _get_title_level(self, para: DocxParagraph) -> int:
        """获取标题层级"""
        try:
            if para.style:
                style_name = para.style.name
                # Title样式是最高级别
                if style_name == 'Title':
                    return 0
                # 从Word标题样式中获取层级
                if style_name.startswith('Heading'):
                    try:
                        return int(style_name.replace('Heading ', ''))
                    except ValueError:
                        pass
        except AttributeError:
            # 如果获取样式失败，继续使用其他方法判断
            pass
                
        # 从文本内容判断层级
        text = para.text.strip()
        for pattern, level in self.title_patterns.items():
            if re.match(pattern, text):
                return level
        return 0